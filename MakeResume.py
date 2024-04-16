import re
import os
import random
from docx import Document
from docx.table import _Cell
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE, WD_BREAK, WD_LINE_SPACING#, WD_STYLE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

# document = Document()

def addParagraph(document,addLine,center,font,spacing):
    paragraph = document.add_paragraph()
    if spacing != None:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(spacing) 
    if center == "center":
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER        
    for item in addLine:
        # print(item)
        run = paragraph.add_run(item.get("text"))
        if font != None:
            run.font.size = Pt(font)
        if item.get("style") == "bold":
            run.bold = True
        elif item.get("style") == "underline":
            run.underline = WD_UNDERLINE.SINGLE  # Single underline (adjust as needed)
        elif item.get("style") == "bullet":
            run.style = document.styles['List Bullet']
        elif item.get("style") == "link":
            run.hyperlink = item["url"]
        elif item.get("style") == "hr":
            paragraph.style.font.size = Pt(12)  # Set font size to minimum to make the line thinner
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # paragraph.add_run().add_break()  # Add a line break to adjust the position
            # paragraph.add_run().add_break()  # Add another line break for better positioning
            # bottom_border = paragraph.add_run()
            
def set_cell_border(cell: _Cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def saveFile(filename, document):
    os.makedirs('/'.join(filename.split('/')[:-1]), exist_ok=True)
    base, ext = os.path.splitext(filename)
    counter = 1
    while os.path.exists(filename):
        filename = f"{base}_{counter}{ext}"
        counter += 1
    document.save(filename)
    return filename

def makeCoverLetter(profile,address,jobDescription,coverLetter):  
    document = Document()  
    addParagraph(document,[{"text": profile[0]}], "center",20,20)
    addParagraph(document,[{"text": "_" * 60, "style":"hr"}], None,10,6)
    addParagraph(document,[{"text": " | ".join(profile[1:4]) + "\n" + " | ".join(profile[4:])}], "center",8,10) 
    addParagraph(document,[{"text": ""}], None,10,10)
    addParagraph(document,[{"text": datetime.today().strftime("%B %d, %Y")}], None,10,10)
    addParagraph(document,[{"text": "\n".join([jobDescription[1]] + address)}], None,10,10)
    addParagraph(document,[{"text": ""}], None,10,10)
    addParagraph(document,[{"text": coverLetter}], None,10,10)
    document.sections[0].top_margin = Inches(0.5) 
    saveFile(f"output/{jobDescription[1]}/CoverLetter-{profile[0].replace(' ','')}.docx", document)

def makeResume(jobDescription,professionalSummary,technicalSkills,skillsProcessed,workExperienceProcessed,expertiseAreas,profile,education):
    document = Document()
    addParagraph(document,[{"text": profile[0]}], "center",20,20)
    addParagraph(document,[{"text": " | ".join(profile[1:4]) + "\n" + " | ".join(profile[4:])}], "center",8,10)
    # addParagraph([{"text": " | ".join(profile[4:])}], "center",8,8)
    addParagraph(document,[{"text": jobDescription[0], "style": "bold"}], "center",14,14)
    technicalSkillsProcessed = []
    for skills in technicalSkills.splitlines():
        skill = skills.replace("- ","").replace("**","").replace("*","")
        if skill.strip() != "":
            technicalSkillsProcessed.append(skill.strip().split(":")[0])
    # print(technicalSkillsProcessed)
    random.shuffle(technicalSkillsProcessed)
    addParagraph(document,[{"text": " | ".join(technicalSkillsProcessed), "style": "bold"}], "center",8,8)
    addParagraph(document,[{"text": professionalSummary}], None,10,10)

    table = document.add_table(rows=len(skillsProcessed), cols=2)
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(6.3) 
    table.style = "Table Grid"
    for index, (key, value) in enumerate(skillsProcessed.items()):
        row = table.rows[index]
        row.cells[0].text = f"{key}:"
        row.cells[1].text = ", ".join(value)
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
        
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(table.columns):
            set_cell_border(
                    table.cell(i, j),
                    top={"sz": 24, "val": "single", "color": "#FFFFFF", "space": "5"},
                    bottom={"sz": 24, "color": "#FFFFFF", "val": "single", "space": "5"},
                    start={"sz": 12, "color": "#FFFFFF", "val": "single"},
                    end={"sz": 12, "color": "#FFFFFF", "val": "single"},
                )
    addParagraph(document,[{"text": ""}], None, 1,1)
    addParagraph(document,[{"text": "Professional Experience", "style": "underline"}], "center", 12,12)
    for job in workExperienceProcessed:
        paragraph = document.add_paragraph()
        jobTitle = paragraph.add_run(job[1])
        jobTitle.font.size = Pt(10)
        jobTitle.bold = True
        company = paragraph.add_run(f" / {job[0]}")
        company.font.size = Pt(10)
        dates = paragraph.add_run(f" ({job[2]}) ")
        dates.font.size = Pt(8)  # Adjust font size as needed
        # jobTitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # company.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT    
        # dates.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for experience in job[3].splitlines():
            match = re.match(r'^[\W_]+', experience)
            if match:
                experience = experience[match.end():]
            # print(experience)
            # addParagraph([{"text": experience, "style": "bullet"}], None, 9,9)
            document.add_paragraph(experience, style='List Bullet').style.font.size = Pt(9)
    addParagraph(document,[{"text": "Education", "style": "underline"}], "center", 12,12) 
    educationProcessed = '\n'.join(education).split(",")    
    addParagraph(document,[{"text": educationProcessed[0], "style":"bold"},{"text": f",{','.join(educationProcessed[1:])}"}], None, 9,10)
    expertiseAreasProcessed = []
    for areas in expertiseAreas.splitlines():
        area = areas.replace("- ","").replace("**","").replace("*","")
        if area.strip() != "":
            expertiseAreasProcessed.append(area.strip().split(":")[0])
    # print(expertiseAreasProcessed)
    addParagraph(document,[{"text": "Areas of Expertise:", "style":"bold"},{"text": f",{', '.join(expertiseAreasProcessed)}"}], None, 9,10)            

    document.sections[0].top_margin = Inches(0.3) 
    document.sections[-1].bottom_margin = Inches(0.3) 
    for section in document.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    saveFile(f"output/{jobDescription[1]}/Resume-{profile[0].replace(' ','')}.docx", document)
    return